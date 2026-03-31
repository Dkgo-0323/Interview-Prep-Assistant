import os
from pathlib import Path
from docx import Document

def setup_fixtures():
    # 1. 确定目录结构
    fixture_dir = Path("tests/fixtures")
    fixture_dir.mkdir(parents=True, exist_ok=True)
    print(f"--- 正在初始化测试集目录: {fixture_dir} ---")

    # 2. 生成正常的简历文档 (含表格)
    def create_sample_resume():
        path = fixture_dir / "sample_resume.docx"
        doc = Document()
        doc.add_heading('个人简历', 0)
        
        doc.add_paragraph('姓名：张三')
        doc.add_paragraph('职位：Python 开发工程师')
        
        # 添加表格测试 _extract_table_text
        doc.add_heading('工作经历', level=1)
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '公司'
        hdr_cells[1].text = '时间'
        hdr_cells[2].text = '职责'
        
        row_cells = table.add_row().cells
        row_cells[0].text = '科技公司 A'
        row_cells[1].text = '2020-2023'
        row_cells[2].text = '后端架构设计与维护'
        
        doc.save(path)
        print(f"已生成: {path.name}")

    # 3. 生成超长文档 (用于测试 MAX_CHARACTERS 截断)
    def create_large_docx():
        path = fixture_dir / "large.docx"
        doc = Document()
        doc.add_heading('超长测试文档', 0)
        
        # 填充超过 50,000 字符
        long_content = "这是重复的测试内容。" * 50  # 每段约 500 字符
        for i in range(120):  # 120 * 500 = 60,000 字符
            doc.add_paragraph(f"段落 {i}: {long_content}")
            
        doc.save(path)
        print(f"已生成: {path.name} (超过 50,000 字符)")

    # 4. 生成损坏的文件 (不是真正的 ZIP 格式)
    def create_corrupted_docx():
        path = fixture_dir / "corrupted.docx"
        with open(path, "wb") as f:
            f.write(b"This is not a zip file, docx must be a zip.")
        print(f"已生成: {path.name} (损坏文件样例)")

    # 执行生成
    create_sample_resume()
    create_large_docx()
    create_corrupted_docx()
    
    print("\n--- 所有测试文件已就绪 ---")
    print("注意: 'encrypted.docx' (加密文件) 建议手动通过 Word 设置密码生成，")
    print("因为 python-docx 不支持直接生成带密码保护的加密文档。")

if __name__ == "__main__":
    setup_fixtures()